import json
import os
from dotenv import load_dotenv
from docx import Document
from google import genai

# 1. Load API Key dari file .env
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    print("Error: API Key tidak ditemukan! Pastikan file .env sudah diisi.")
    exit()

# 2. Setup Gemini API
client = genai.Client(api_key=api_key)

# 3. DAFTAR KANDIDAT MODEL (Auto-Fallback)
kandidat_model = [
    'gemini-2.5-flash',
    'gemini-2.0-flash',
    'gemini-1.5-flash-latest',
    'gemini-1.5-flash',
    'gemini-1.5-pro',
    'gemini-1.0-pro',
    'gemini-flash-latest'
]

model_aktif = None

print("=== MENCARI MODEL AI YANG TERSEDIA ===")
for nama_model in kandidat_model:
    try:
        print(f"[*] Mengetes model: {nama_model}...")
        # Lakukan tes ringan untuk melihat apakah model ini bisa diakses
        tes_koneksi = client.models.generate_content(
            model=nama_model,
            contents="Halo, balas dengan kata 'OK' jika kamu aktif."
        )
        if tes_koneksi.text:
            model_aktif = nama_model
            print(f"✅ BERHASIL! Kita akan menggunakan model: {model_aktif}\n")
            break
    except Exception as e:
        print(f"❌ Gagal: {nama_model} tidak tersedia.")

# Jika semua model di list gagal
if not model_aktif:
    print("\n[!] ERROR FATAL: Tidak ada satupun model yang bisa diakses.")
    print("Kemungkinan penyebab: API Key limit habis, atau region API tidak didukung.")
    exit()

# 4. Persiapan Folder Output
folder_output = "Hasil_Novel"
if not os.path.exists(folder_output):
    os.makedirs(folder_output)
    print(f"Folder '{folder_output}' siap.")

# 5. Baca Blueprint JSON
print("Membaca dbase.json...")
try:
    with open('dbase.json', 'r', encoding='utf-8') as file:
        dbase = json.load(file)
except FileNotFoundError:
    print("Error: File dbase.json tidak ditemukan!")
    exit()

info_dasar = f"""
Judul: {dbase['metadata']['judul_sementara']}
Genre: {dbase['metadata']['genre']}
Gaya Bahasa: {dbase['metadata']['gaya_bahasa']}
Latar Tempat: {', '.join(dbase['dunia_dan_latar']['tempat_utama'])}
Karakter: {dbase['karakter']}
"""

ringkasan_sebelumnya = "Ini adalah bab pertama. Mulai cerita dari awal sesuai plot dan langsung bangun ketegangan."

# 6. Loop Eksekusi per Bab
for bab in dbase['bab']:
    print(f"\n[+] Sedang memproduksi Episode {bab['bab_ke']}: {bab['judul_bab']}...")
    
    prompt_nulis = f"""
    Kamu adalah penulis novel misteri profesional dari Indonesia. Tuliskan Episode {bab['bab_ke']} dengan judul '{bab['judul_bab']}'.
    
    INFORMASI CERITA:
    {info_dasar}
    
    RINGKASAN EPISODE SEBELUMNYA (Lanjutkan cerita dengan mulus dari titik ini, perhatikan kontinuitas waktu dan posisi tokoh):
    {ringkasan_sebelumnya}
    
    PLOT EPISODE INI:
    - Apa yang terjadi: {bab['plot_utama']}
    - Konflik: {bab['konflik_bab_ini']}
    - Akhir episode: {bab['ending_bab']}
    
    INSTRUKSI PENULISAN:
    - Tulis minimal 1500 kata.
    - Jangan berikan judul bab di dalam teks cerita, langsung mulai ke isi cerita.
    - Format teks agar rapi (ada paragrafnya).
    - Jangan pernah membuat ringkasan di bagian ini.
    """
    
    try:
        # Generate isi novel dengan model yang terbukti berhasil
        response_bab = client.models.generate_content(
            model=model_aktif,
            contents=prompt_nulis
        )
        isi_bab = response_bab.text
        
        # Bikin File DOCX
        doc = Document()
        doc.add_heading(f"Episode {bab['bab_ke']}: {bab['judul_bab']}", level=1)
        doc.add_paragraph(isi_bab)
        
        nama_file = f"Episode_{bab['bab_ke']:02d}_{bab['judul_bab'].replace(' ', '_')}.docx"
        path_file = os.path.join(folder_output, nama_file)
        
        # Simpan file DOCX
        doc.save(path_file)
        print(f"    -> Berhasil disimpan di folder {folder_output} sebagai {nama_file}")

        # Generate Ringkasan untuk next episode
        print("    -> Meracik ringkasan untuk kontinuitas ke episode besok...")
        prompt_summary = f"""
        Buat ringkasan padat (maksimal 2 paragraf) dari teks novel berikut. 
        Fokus pada posisi dan kondisi terakhir karakter, serta detail menggantung (cliffhanger).
        Teks novel: {isi_bab}
        """
        response_summary = client.models.generate_content(
            model=model_aktif,
            contents=prompt_summary
        )
        ringkasan_sebelumnya = response_summary.text
        
    except Exception as e:
        print(f"Error di Episode {bab['bab_ke']}: {e}")
        break
    
print(f"\nSelesai Bosku! Semua file Word sudah rapi di dalam folder '{folder_output}'.")