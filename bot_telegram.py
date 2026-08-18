import json
import os
import telebot
from dotenv import load_dotenv
from docx import Document
from google import genai

# 1. Load Keys
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
tele_token = os.getenv("TELEGRAM_TOKEN")

if not api_key or not tele_token:
    print("Error: API Key Gemini atau Token Telegram tidak ditemukan di .env!")
    exit()

bot = telebot.TeleBot(tele_token)
client = genai.Client(api_key=api_key)

# 2. Setup Auto-Fallback Gemini
kandidat_model = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash-latest', 'gemini-1.5-flash']
model_aktif = None
print("Mencari model AI...")
for nama_model in kandidat_model:
    try:
        if client.models.generate_content(model=nama_model, contents="tes").text:
            model_aktif = nama_model
            print(f"✅ Model AI siap: {model_aktif}")
            break
    except: continue

if not model_aktif:
    print("[!] Gagal terhubung ke Google AI.")
    exit()

# Folder & Memori Setup
folder_output = "Hasil_Novel"
file_memori = "ingatan_ai.txt"
if not os.path.exists(folder_output): os.makedirs(folder_output)

# 3. Pesan Selamat Datang (Start)
@bot.message_handler(commands=['start', 'help'])
def send_welcome(message):
    teks = """Halo Bosku! ☕
Bot Penulis Novel siap menerima perintah.

Ketik perintah dengan format berikut, pisahkan dengan garis lurus ( | )
`/tulis <Bab_ke> | <Judul> | <Plot_Utama> | <Konflik> | <Ending>`

Contoh:
`/tulis 4 | Balik Kandang | Agus kabur ke kandang | Ayam ribut | HP Agus hilang`
"""
    bot.reply_to(message, teks, parse_mode='Markdown')

# 4. Mesin Utama (Menerima Perintah /tulis)
@bot.message_handler(commands=['tulis'])
def handle_tulis(message):
    try:
        # Parsing teks dari Telegram
        raw_text = message.text.replace("/tulis", "").strip()
        parts = [p.strip() for p in raw_text.split("|")]
        
        if len(parts) != 5:
            bot.reply_to(message, "❌ Format salah Bosku! Pastikan ada 5 bagian dipisah tanda |")
            return
            
        bab_ke, judul, plot, konflik, ending = parts
        bab_ke = int(bab_ke)
        
        bot.reply_to(message, f"Siap Bosku! 🚀 Sedang meracik *Episode {bab_ke}: {judul}*... Sambil seruput kopinya dulu.", parse_mode='Markdown')

        # Load Database & Info Dasar
        with open('dbase.json', 'r', encoding='utf-8') as f: dbase = json.load(f)
        info_dasar = f"Judul: {dbase['metadata']['judul_sementara']}\nGenre: {dbase['metadata']['genre']}\nGaya: {dbase['metadata']['gaya_bahasa']}\nKarakter: {dbase['karakter']}"

        # Load Memori
        ringkasan_sebelumnya = "Ini bab pertama. Langsung bangun ketegangan."
        if os.path.exists(file_memori):
            with open(file_memori, 'r', encoding='utf-8') as f: ringkasan_sebelumnya = f.read()

        # Generate Cerita
        prompt_nulis = f"""
        Tuliskan Episode {bab_ke} dengan judul '{judul}'.
        INFORMASI CERITA: {info_dasar}
        RINGKASAN EPISODE SEBELUMNYA: {ringkasan_sebelumnya}
        PLOT EPISODE INI:
        - Kejadian: {plot}
        - Konflik: {konflik}
        - Akhir: {ending}
        
        INSTRUKSI: Tulis novel profesional, min 800 kata. Tanpa judul di dalam teks. Jangan buat ringkasan.
        """
        response_bab = client.models.generate_content(model=model_aktif, contents=prompt_nulis)
        isi_bab = response_bab.text
        
        # Simpan ke Word
        doc = Document()
        doc.add_heading(f"Episode {bab_ke}: {judul}", level=1)
        doc.add_paragraph(isi_bab)
        
        nama_file = f"Episode_{bab_ke:02d}_{judul.replace(' ', '_')}.docx"
        path_file = os.path.join(folder_output, nama_file)
        doc.save(path_file)

        # Update Memori (Summary)
        bot.send_message(message.chat.id, "✅ Cerita beres! Lagi nyiapin memori buat besok...")
        prompt_summary = f"Buat ringkasan padat dari novel berikut. Fokus posisi terakhir karakter dan cliffhanger. Teks: {isi_bab}"
        response_summary = client.models.generate_content(model=model_aktif, contents=prompt_summary)
        
        with open(file_memori, 'w', encoding='utf-8') as f:
            f.write(response_summary.text)

        # Update dbase.json agar data tidak hilang
        bab_baru = {
            "bab_ke": bab_ke,
            "judul_bab": judul,
            "plot_utama": plot,
            "konflik_bab_ini": konflik,
            "ending_bab": ending
        }
        # Cek apakah bab sudah ada di json, kalau belum tambahkan
        if not any(b['bab_ke'] == bab_ke for b in dbase['bab']):
            dbase['bab'].append(bab_baru)
            with open('dbase.json', 'w', encoding='utf-8') as f:
                json.dump(dbase, f, indent=2)

        # Kirim File Word ke Telegram
        with open(path_file, 'rb') as doc_file:
            bot.send_document(message.chat.id, doc_file, caption=f"Done Bosku! Ini file {nama_file}")

    except Exception as e:
        bot.reply_to(message, f"❌ Waduh, ada error Bosku:\n{str(e)}")

print("\n[🤖] Bot Telegram aktif. Silakan chat bot dari HP Bosku!")
# Jalankan bot terus menerus
bot.infinity_polling()