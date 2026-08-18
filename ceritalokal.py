import json
import os
from dotenv import load_dotenv
from docx import Document
from google import genai
import telebot

# 1. Load API Keys (dari Environment Variables / GitHub Secrets)
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
tele_token = os.getenv("TELEGRAM_TOKEN")
chat_id = os.getenv("TELEGRAM_CHAT_ID") # ID Chat Telegram Bosku

if not api_key:
    print("Error: GEMINI_API_KEY tidak ditemukan!")
    exit()

client = genai.Client(api_key=api_key)

# Setup Bot Telegram jika token & chat_id tersedia (Opsional untuk notifikasi)
bot = telebot.TeleBot(tele_token) if tele_token else None

# 2. Auto-Fallback Model Gemini
kandidat_model = ['gemini-2.5-flash', 'gemini-2.0-flash', 'gemini-1.5-flash-latest', 'gemini-1.5-flash']
model_aktif = None
for nama_model in kandidat_model:
    try:
        if client.models.generate_content(model=nama_model, contents="tes").text:
            model_aktif = nama_model
            break
    except: continue

if not model_aktif:
    print("[!] Model AI gagal terhubung.")
    exit()

# 3. Setup Folder & File Memori
folder_output = "Hasil_Novel"
file_memori = "ingatan_ai.txt"
file_db = "dbase.json"

if not os.path.exists(folder_output): os.makedirs(folder_output)

# Read JSON
with open(file_db, 'r', encoding='utf-8') as f:
    dbase = json.load(f)

# Tentukan Episode Selanjutnya
episode_terakhir = max([b['bab_ke'] for b in dbase['bab']]) if dbase['bab'] else 0
episode_baru = episode_terakhir + 1

# Read Memori
ringkasan_sebelumnya = "Ini bab pertama. Mulai cerita dari awal."
if os.path.exists(file_memori):
    with open(file_memori, 'r', encoding='utf-8') as f:
        ringkasan_sebelumnya = f.read()

print(f"🤖 Memulai auto-generate untuk Episode {episode_baru}...")

# 4. TAHAP 1: AI MENULIS PLOT OTOMATIS
prompt_bikin_plot = f"""
Kamu adalah Sutradara / Penulis Skenario Novel Misteri Thriller Lokal Indonesia.
Judul Novel: {dbase['metadata']['judul_sementara']}
Genre: {dbase['metadata']['genre']}
Karakter: {dbase['karakter']}

RINGKASAN EPISODE TERAKHIR:
{ringkasan_sebelumnya}

Buatkan outline/plot singkat untuk Episode {episode_baru} yang melanjutkan cerita secara logis, seru, dan penuh ketegangan!
Kembalikan hasilnya HANYA dalam format JSON valid berikut (tanpa format markdown tambahan/tanpa backtick json):
{{
  "judul_bab": "Judul Bab Yang Menarik",
  "plot_utama": "Penjelasan ringkas kejadian di bab ini",
  "konflik_bab_ini": "Konflik atau rintangan yang dihadapi",
  "ending_bab": "Ending menggantung / cliffhanger"
}}
"""

try:
    res_plot = client.models.generate_content(model=model_aktif, contents=prompt_bikin_plot)
    raw_plot_text = res_plot.text.strip().replace("```json", "").replace("```", "").strip()
    plot_json = json.loads(raw_plot_text)
    print(f"✅ Plot Episode {episode_baru} berhasil dibuat: {plot_json['judul_bab']}")
except Exception as e:
    print(f"❌ Gagal membuat plot otomatis: {e}")
    exit()

# 5. TAHAP 2: AI MENULIS NOVEL UTUH
info_dasar = f"Judul: {dbase['metadata']['judul_sementara']}\nGenre: {dbase['metadata']['genre']}\nGaya: {dbase['metadata']['gaya_bahasa']}"

prompt_nulis = f"""
Kamu adalah penulis novel misteri profesional. Tuliskan Episode {episode_baru} dengan judul '{plot_json['judul_bab']}'.

INFORMASI CERITA:
{info_dasar}

RINGKASAN EPISODE SEBELUMNYA:
{ringkasan_sebelumnya}

PLOT EPISODE INI:
- Kejadian: {plot_json['plot_utama']}
- Konflik: {plot_json['konflik_bab_ini']}
- Ending: {plot_json['ending_bab']}

INSTRUKSI:
- Tulis minimal 800 kata. Format teks rapi berparagraf.
- Jangan tulis judul di dalam teks cerita.
- Jangan buat ringkasan di bagian ini.
"""

res_cerita = client.models.generate_content(model=model_aktif, contents=prompt_nulis)
isi_bab = res_cerita.text

# 6. SIMPAN FILE WORD
doc = Document()
doc.add_heading(f"Episode {episode_baru}: {plot_json['judul_bab']}", level=1)
doc.add_paragraph(isi_bab)

nama_file = f"Episode_{episode_baru:02d}_{plot_json['judul_bab'].replace(' ', '_')}.docx"
path_file = os.path.join(folder_output, nama_file)
doc.save(path_file)
print(f"✅ File Word berhasil disimpan: {path_file}")

# 7. UPDATE MEMORI UNTUK EPISODE BESOK
prompt_summary = f"Buat ringkasan padat dari novel berikut. Fokus posisi terakhir karakter dan cliffhanger. Teks: {isi_bab}"
res_summary = client.models.generate_content(model=model_aktif, contents=prompt_summary)

with open(file_memori, 'w', encoding='utf-8') as f:
    f.write(res_summary.text)

# 8. UPDATE DATABASE JSON
data_bab_baru = {
    "bab_ke": episode_baru,
    "judul_bab": plot_json['judul_bab'],
    "plot_utama": plot_json['plot_utama'],
    "konflik_bab_ini": plot_json['konflik_bab_ini'],
    "ending_bab": plot_json['ending_bab']
}
dbase['bab'].append(data_bab_baru)
with open(file_db, 'w', encoding='utf-8') as f:
    json.dump(dbase, f, indent=2)

# 9. KIRIM NOTIFIKASI TELEGRAM (OPTIONAL)
if bot and chat_id:
    try:
        with open(path_file, 'rb') as doc_file:
            bot.send_document(
                chat_id, 
                doc_file, 
                caption=f"🚀 *Auto-Generate Berhasil!*\n\nEpisode {episode_baru}: *{plot_json['judul_bab']}*\nFile Word sudah terbit & tersimpan di GitHub!",
                parse_mode='Markdown'
            )
        print("✅ Notifikasi terkirim ke Telegram!")
    except Exception as e:
        print(f"⚠️ Gagal kirim ke Telegram: {e}")

print("🎉 Selesai 100%! Siap di-commit otomatis oleh GitHub.")